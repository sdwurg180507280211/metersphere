#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成全流程平台对接功能需求说明书 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ========== 全局样式 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    heading_style.font.name = '黑体'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)


def add_para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('全流程平台对接功能需求说明书')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('——需求平台侧开发需求')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 0)

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    '文档版本：V1.0',
    '编制日期：2026年5月7日',
    '编制方：测试平台团队',
    '交付方：需求平台团队',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)

doc.add_page_break()

# ========== 目录页 ==========
doc.add_heading('目录', level=1)
toc_items = [
    '1  引言',
    '2  项目介绍',
    '3  项目背景',
    '4  开发目标',
    '5  对接流程',
    '  5.1  总体架构',
    '  5.2  需求平台需实现的功能',
    '  5.3  消息推送规范',
    '  5.4  状态回传接收规范',
    '  5.5  接口安全与可靠性',
    '6  最终效果',
    '  6.1  需求推送效果',
    '  6.2  测试状态回传效果',
    '  6.3  端到端闭环效果',
    '7  附录',
    '  7.1  需求平台推送字段清单',
    '  7.2  测试平台回传字段清单',
    '  7.3  消息样例',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ========== 1 引言 ==========
doc.add_heading('1  引言', level=1)

doc.add_paragraph(
    '本文档由测试平台团队编制，面向需求平台开发团队，旨在明确全流程平台对接中需求平台侧需要完成的开发工作。'
    '文档详细描述了需求平台应实现的功能、消息推送规范、回传消息接收规范及联调验收标准，'
    '作为需求平台侧开发的依据和双方联调的参考。'
)

doc.add_paragraph(
    '本文档不涉及测试平台（MeterSphere）内部的设计与实现细节，仅从对接视角描述需求平台需交付的内容。'
)

# ========== 2 项目介绍 ==========
doc.add_heading('2  项目介绍', level=1)

doc.add_paragraph(
    '本项目实现需求平台与MeterSphere测试平台的双向数据对接，通过RocketMQ消息队列建立异步通信通道。'
    '核心目标是打通"需求→测试"的数据链路，实现需求驱动的测试管理闭环。'
)

add_table(
    ['项目', '说明'],
    [
        ['项目名称', '全流程平台对接（需求平台 ↔ MeterSphere测试平台）'],
        ['对接方式', 'RocketMQ消息队列，双向异步通信'],
        ['消息方向1', '需求平台 → 测试平台：需求同步消息'],
        ['消息方向2', '测试平台 → 需求平台：测试状态回传消息'],
        ['版本范围', 'V1 最小可交付版本'],
        ['项目周期', '2026年3月10日 - 2026年6月30日'],
    ]
)

# ========== 3 项目背景 ==========
doc.add_heading('3  项目背景', level=1)

doc.add_paragraph(
    '当前需求平台与测试平台之间缺乏数据通道，测试人员需要手工从需求平台抄写需求信息到测试平台，'
    '测试进度也无法自动反馈给需求平台。这导致以下问题：'
)

problems = [
    '信息孤岛：需求与测试数据分离，无法自动关联和追溯',
    '重复录入：测试人员需手工录入需求信息，效率低且易出错',
    '状态脱节：需求平台无法实时感知测试进度，影响项目决策',
    '追溯困难：需求到测试的链路断裂，无法量化测试覆盖',
]
for p_text in problems:
    doc.add_paragraph(p_text, style='List Bullet')

doc.add_paragraph(
    '通过建立双向数据通道，需求平台自动推送需求数据，测试平台自动回传测试状态，'
    '实现需求全生命周期的可追溯管理。'
)

# ========== 4 开发目标 ==========
doc.add_heading('4  开发目标', level=1)

doc.add_heading('4.1  功能目标', level=2)

goals = [
    '需求平台能够将需求信息（新建/更新/取消）通过RocketMQ实时推送到测试平台',
    '需求平台能够接收测试平台回传的测试计划状态信息',
    '需求平台在收到回传消息后，能够在需求详情中展示测试进度和报告链接',
    '双方通过traceId实现全链路追踪，便于问题排查',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('4.2  性能目标', level=2)

add_table(
    ['指标', '目标值'],
    [
        ['消息推送延迟', '< 5秒（从需求变更到消息发出）'],
        ['消息消费延迟', '< 5秒（测试平台从收到消息到落库）'],
        ['回传消息消费延迟', '< 5秒（需求平台从收到回传到更新状态）'],
        ['消息可靠性', '不丢失、不重复处理（幂等保障）'],
    ]
)

doc.add_heading('4.3  安全目标', level=2)

security = [
    'RocketMQ启用ACL权限控制（AccessKey/SecretKey）',
    'Topic级别读写权限隔离',
    '消息传输使用TLS/SSL加密通道',
    '敏感字段按需加密',
]
for s in security:
    doc.add_paragraph(s, style='List Bullet')

# ========== 5 对接流程 ==========
doc.add_heading('5  对接流程', level=1)

doc.add_heading('5.1  总体架构', level=2)

doc.add_paragraph(
    '双方通过RocketMQ消息队列进行双向异步通信，架构如下：'
)

arch_text = (
    '┌─────────────────┐         RocketMQ          ┌──────────────────┐\n'
    '│   需求平台      │ ──────────────────────────>│  测试平台        │\n'
    '│                 │  topic-requirement-to-ms   │  (MeterSphere)   │\n'
    '│                 │                            │                  │\n'
    '│                 │  topic-ms-to-requirement   │                  │\n'
    '│                 │ <──────────────────────────│                  │\n'
    '└─────────────────┘         RocketMQ          └──────────────────┘'
)
p = doc.add_paragraph()
run = p.add_run(arch_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph(
    '核心流程：\n'
    '1. 需求平台通过RocketMQ发送需求消息（新建/更新/取消）\n'
    '2. 测试平台消费消息，写入需求池\n'
    '3. 测试人员在测试平台创建测试计划\n'
    '4. 测试计划状态变化时，测试平台通过RocketMQ回传给需求平台\n'
    '5. 需求平台消费回传消息，更新需求测试状态'
)

doc.add_heading('5.2  需求平台需实现的功能', level=2)

doc.add_heading('5.2.1  需求同步消息推送', level=3)

doc.add_paragraph(
    '需求平台在需求状态变更时，主动推送消息到RocketMQ。这是对接的核心前提，'
    '需求平台需要在以下场景触发消息推送：'
)

add_table(
    ['触发场景', 'operationType', '说明'],
    [
        ['新建需求', 'CREATED', '需求首次创建时推送，包含完整需求信息'],
        ['更新需求', 'UPDATED', '需求信息变更时推送，包含最新完整需求信息'],
        ['取消需求', 'CANCELLED', '需求被取消/撤回时推送，仅需要dmpNum'],
    ]
)

doc.add_paragraph('')
add_para('推送要求：', bold=True)
requirements = [
    '消息格式为JSON，UTF-8编码',
    '每条消息必须包含 dmpNum（需求编号）和 operationType（操作类型）',
    '每条消息必须包含 eventTime（消息事件时间戳），用于乱序消息判断',
    '每条消息建议携带 traceId（追踪ID），便于全链路问题排查',
    '推送失败时需有重试机制（RocketMQ Producer 侧重试）',
    '确保消息不丢失（建议使用同步发送或事务消息）',
]
for r in requirements:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('5.2.2  测试状态回传消息消费', level=3)

doc.add_paragraph(
    '需求平台需要消费测试平台回传的消息，获取测试计划状态和报告链接。'
    '收到回传消息后，需求平台应：'
)

callback_actions = [
    '根据 dmpNum 找到对应需求',
    '更新需求的测试状态（planStatus）',
    '记录测试计划的计划起止时间、实际起止时间',
    '记录测试负责人',
    '记录测试报告分享链接（planShareUrl），并在需求详情页提供可点击的链接',
    '基于 syncTime 和 traceId 记录回传日志',
]
for a in callback_actions:
    doc.add_paragraph(a, style='List Bullet')

doc.add_paragraph('')
add_para('消费要求：', bold=True)
consume_requirements = [
    '消费失败时不要确认消息，让RocketMQ重试',
    '实现幂等消费：同一 dmpNum + syncTime 的消息不重复处理',
    '回传消息不影响需求平台主业务流程，消费异常应记录日志而非阻塞',
]
for r in consume_requirements:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('5.2.3  RocketMQ环境准备', level=3)

doc.add_paragraph('需求平台需要准备以下RocketMQ环境：')

add_table(
    ['配置项', '要求', '说明'],
    [
        ['NameServer地址', '需与测试平台网络互通', '双方连接同一个RocketMQ集群'],
        ['Topic: 需求同步', 'topic-requirement-to-metersphere', '需求平台为Producer，测试平台为Consumer'],
        ['Topic: 状态回传', 'topic-metersphere-to-requirement', '测试平台为Producer，需求平台为Consumer'],
        ['Producer Group（需求→测试）', 'producer-requirement-to-metersphere', '需求平台发送同步消息的Producer Group'],
        ['Consumer Group（需求→测试）', 'consumer-requirement-to-metersphere', '测试平台消费同步消息的Consumer Group'],
        ['Producer Group（测试→需求）', 'producer-metersphere-to-requirement', '测试平台发送回传消息的Producer Group'],
        ['Consumer Group（测试→需求）', 'consumer-metersphere-to-requirement', '需求平台消费回传消息的Consumer Group'],
        ['ACL权限', 'AccessKey/SecretKey', '需求平台需要有上述两个Topic的读写权限'],
    ]
)

doc.add_heading('5.3  消息推送规范', level=2)

doc.add_heading('5.3.1  需求同步消息格式', level=3)

doc.add_paragraph('Topic: topic-requirement-to-metersphere')
doc.add_paragraph('消息体JSON格式如下：')

msg_fields = [
    ('dmpNum', 'String', '是', '需求编号，唯一标识'),
    ('name1', 'String', '是', '需求名称'),
    ('operationType', 'String', '是', '操作类型：CREATED / UPDATED / CANCELLED'),
    ('reqManagerName', 'String', '否', '需求负责人'),
    ('actName', 'String', '否', '当前环节（流程节点名称）'),
    ('createTime', 'Long', '否', '需求提出时间（毫秒时间戳）'),
    ('parentWfinstCode', 'String', '否', '主流程编码'),
    ('reqFatherClass', 'String', '否', '需求大类'),
    ('reqSonClass', 'String', '否', '需求子类'),
    ('systemName', 'String', '否', '所属系统'),
    ('upTime', 'Long', '否', '预计上线时间（毫秒时间戳）'),
    ('assigneeName', 'String', '否', '当前处理人'),
    ('createdept', 'String', '否', '需求申请部门（注意：字段名无下划线）'),
    ('createUser1', 'String', '否', '需求申请人（注意：字段名后缀为1）'),
    ('deptName', 'String', '否', '需求负责人处室'),
    ('startUserName', 'String', '否', '创建人'),
    ('eventTime', 'Long', '是', '消息事件时间（毫秒时间戳），用于幂等和乱序判断'),
    ('traceId', 'String', '否', '追踪ID，全链路问题排查'),
]

add_table(
    ['字段名', '类型', '必填', '说明'],
    [(f[0], f[1], f[2], f[3]) for f in msg_fields]
)

doc.add_paragraph('')
add_para('字段映射注意事项：', bold=True)
notes = [
    'name1 → 需求名称（需求平台字段名为name1，非name）',
    'createdept → 需求申请部门（需求平台字段名无下划线，非createdDept）',
    'createUser1 → 需求申请人（注意后缀为1，非createUser）',
    'operationType 与 actName 是不同概念：operationType是消息操作类型（CREATED/UPDATED/CANCELLED），actName是流程环节名称（如"测试待处理"）',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_heading('5.3.2  operationType说明', level=3)

add_table(
    ['operationType', '含义', '触发时机', '必填字段'],
    [
        ['CREATED', '新建需求', '需求首次创建', 'dmpNum, name1, eventTime'],
        ['UPDATED', '更新需求', '需求信息变更', 'dmpNum, name1, eventTime'],
        ['CANCELLED', '取消需求', '需求被取消/撤回', 'dmpNum, eventTime'],
    ]
)

doc.add_heading('5.4  状态回传接收规范', level=2)

doc.add_heading('5.4.1  回传消息格式', level=3)

doc.add_paragraph('Topic: topic-metersphere-to-requirement')
doc.add_paragraph('测试平台回传消息体JSON格式如下：')

callback_fields = [
    ('dmpNum', 'String', '需求编号（关联主键）'),
    ('planStatus', 'String', '测试计划状态'),
    ('plannedStartTime', 'Long', '计划开始时间（毫秒时间戳）'),
    ('plannedEndTime', 'Long', '计划结束时间（毫秒时间戳）'),
    ('actualStartTime', 'Long', '实际开始时间（毫秒时间戳）'),
    ('actualEndTime', 'Long', '实际结束时间（毫秒时间戳）'),
    ('principalUsers', 'String', '测试负责人'),
    ('planShareUrl', 'String', '测试报告分享链接'),
    ('syncTime', 'Long', '同步时间（毫秒时间戳）'),
    ('traceId', 'String', '追踪ID'),
]

add_table(
    ['字段名', '类型', '说明'],
    [(f[0], f[1], f[2]) for f in callback_fields]
)

doc.add_paragraph('')
add_para('planStatus状态值说明：', bold=True)

add_table(
    ['planStatus', '含义', '说明'],
    [
        ['PENDING', '未开始', '测试计划已创建，尚未开始执行'],
        ['PREPARED', '已准备', '测试计划已完成准备'],
        ['RUNNING', '进行中', '测试计划正在执行'],
        ['COMPLETED', '已完成', '测试计划已完成执行'],
        ['FINISHED', '已结束', '测试计划已结束'],
    ]
)

doc.add_paragraph('')
add_para('planShareUrl格式说明：', bold=True)
doc.add_paragraph(
    '链接格式为 /track/share-plan-report?shareId={shareId}，'
    '需求平台拼接测试平台域名前缀即可直接访问。例如：https://metersphere.example.com/track/share-plan-report?shareId=abc123'
)

doc.add_heading('5.5  接口安全与可靠性', level=2)

doc.add_heading('5.5.1  消息可靠性', level=3)

reliability = [
    '需求平台发送消息时使用同步发送（syncSend）或事务消息，确保消息不丢失',
    '消息消费失败时，由RocketMQ自动重试（默认16次），不要手动确认失败消息',
    '双方实现幂等消费：基于 dmpNum + eventTime/syncTime 判断是否重复',
    '乱序消息处理：收到 eventTime 早于已处理记录的消息时，直接丢弃',
]
for r in reliability:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('5.5.2  数据安全', level=3)

data_security = [
    'RocketMQ启用ACL权限控制，需求平台和测试平台使用不同的AccessKey',
    'Topic级别权限隔离：需求平台仅对同步Topic有写权限、对回传Topic有读权限',
    '建议使用TLS/SSL加密通道传输',
    'traceId贯穿全链路，便于审计追溯',
]
for d in data_security:
    doc.add_paragraph(d, style='List Bullet')

# ========== 6 最终效果 ==========
doc.add_heading('6  最终效果', level=1)

doc.add_heading('6.1  需求推送效果', level=2)

doc.add_paragraph(
    '需求平台在需求创建、更新、取消时，自动通过RocketMQ推送消息。'
    '测试平台实时消费消息，需求信息在测试平台的需求池中立即可见。'
)

push_effects = [
    '新建需求后，测试平台5秒内可在需求池列表看到该需求',
    '更新需求后，测试平台的需求池信息同步更新',
    '取消需求后，测试平台需求池状态自动变为"已取消"，禁止创建测试计划',
]
for e in push_effects:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('6.2  测试状态回传效果', level=2)

doc.add_paragraph(
    '测试平台在测试计划状态变更时，自动通过RocketMQ回传消息。'
    '需求平台消费回传消息后，在需求详情中展示测试进度。'
)

callback_effects = [
    '需求详情页展示测试计划状态（未开始/进行中/已完成等）',
    '需求详情页展示测试负责人',
    '需求详情页展示测试计划起止时间和实际起止时间',
    '需求详情页提供可点击的测试报告链接，直接跳转查看报告',
    '需求列表页可按测试状态筛选需求',
]
for e in callback_effects:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('6.3  端到端闭环效果', level=2)

doc.add_paragraph('对接完成后的完整数据流转：')

flow_text = (
    '1. 需求平台创建需求\n'
    '    ↓  RocketMQ推送 CREATED 消息\n'
    '2. 测试平台需求池自动展示新需求（状态：未创建）\n'
    '    ↓  测试人员点击"创建测试计划"\n'
    '3. 测试平台创建测试计划（状态：未开始）\n'
    '    ↓  RocketMQ回传 planStatus=PENDING\n'
    '4. 需求平台展示"测试未开始"\n'
    '    ↓  测试人员开始执行\n'
    '5. 测试平台测试计划状态变为"进行中"\n'
    '    ↓  RocketMQ回传 planStatus=RUNNING\n'
    '6. 需求平台展示"测试进行中"\n'
    '    ↓  测试执行完成\n'
    '7. 测试平台测试计划状态变为"已完成"\n'
    '    ↓  RocketMQ回传 planStatus=COMPLETED + 报告链接\n'
    '8. 需求平台展示"测试已完成"，提供报告链接'
)
p = doc.add_paragraph()
run = p.add_run(flow_text)
run.font.size = Pt(10)

# ========== 7 附录 ==========
doc.add_heading('7  附录', level=1)

doc.add_heading('7.1  需求平台推送字段清单', level=2)

add_table(
    ['序号', '字段名', '类型', '必填', '说明'],
    [
        ['1', 'dmpNum', 'String', '是', '需求编号（唯一主键）'],
        ['2', 'name1', 'String', '是', '需求名称'],
        ['3', 'operationType', 'String', '是', '操作类型：CREATED/UPDATED/CANCELLED'],
        ['4', 'eventTime', 'Long', '是', '消息事件时间戳（毫秒）'],
        ['5', 'reqManagerName', 'String', '否', '需求负责人'],
        ['6', 'actName', 'String', '否', '当前环节'],
        ['7', 'createTime', 'Long', '否', '需求提出时间戳（毫秒）'],
        ['8', 'parentWfinstCode', 'String', '否', '主流程编码'],
        ['9', 'reqFatherClass', 'String', '否', '需求大类'],
        ['10', 'reqSonClass', 'String', '否', '需求子类'],
        ['11', 'systemName', 'String', '否', '所属系统'],
        ['12', 'upTime', 'Long', '否', '预计上线时间戳（毫秒）'],
        ['13', 'assigneeName', 'String', '否', '当前处理人'],
        ['14', 'createdept', 'String', '否', '需求申请部门'],
        ['15', 'createUser1', 'String', '否', '需求申请人'],
        ['16', 'deptName', 'String', '否', '需求负责人处室'],
        ['17', 'startUserName', 'String', '否', '创建人'],
        ['18', 'traceId', 'String', '否', '追踪ID'],
    ]
)

doc.add_heading('7.2  测试平台回传字段清单', level=2)

add_table(
    ['序号', '字段名', '类型', '说明'],
    [
        ['1', 'dmpNum', 'String', '需求编号（关联主键）'],
        ['2', 'planStatus', 'String', '测试计划状态'],
        ['3', 'plannedStartTime', 'Long', '计划开始时间戳（毫秒）'],
        ['4', 'plannedEndTime', 'Long', '计划结束时间戳（毫秒）'],
        ['5', 'actualStartTime', 'Long', '实际开始时间戳（毫秒）'],
        ['6', 'actualEndTime', 'Long', '实际结束时间戳（毫秒）'],
        ['7', 'principalUsers', 'String', '测试负责人'],
        ['8', 'planShareUrl', 'String', '测试报告分享链接'],
        ['9', 'syncTime', 'Long', '同步时间戳（毫秒）'],
        ['10', 'traceId', 'String', '追踪ID'],
    ]
)

doc.add_heading('7.3  消息样例', level=2)

doc.add_heading('7.3.1  需求同步消息样例', level=3)

doc.add_paragraph('Topic: topic-requirement-to-metersphere')

sample1 = '''{
  "dmpNum": "ICBS-POS-20260424-7371",
  "name1": "银保微投新产品-瑞众护身甲意外保险需求",
  "operationType": "CREATED",
  "reqManagerName": "张三",
  "actName": "测试待处理",
  "createTime": 1713945600000,
  "parentWfinstCode": "WF-20260424-001",
  "reqFatherClass": "功能需求",
  "reqSonClass": "新增功能",
  "systemName": "瑞众保险个险核心业务系统-保全",
  "upTime": 1718294400000,
  "assigneeName": "李四",
  "createdept": "产品部",
  "createUser1": "王五",
  "deptName": "产品一部",
  "startUserName": "赵六",
  "eventTime": 1713945600000,
  "traceId": "trace-20260424-7371-001"
}'''
p = doc.add_paragraph()
run = p.add_run(sample1)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('7.3.2  状态回传消息样例', level=3)

doc.add_paragraph('Topic: topic-metersphere-to-requirement')

sample2 = '''{
  "dmpNum": "ICBS-POS-20260424-7371",
  "planStatus": "COMPLETED",
  "plannedStartTime": 1714560000000,
  "plannedEndTime": 1716230400000,
  "actualStartTime": 1714646400000,
  "actualEndTime": 1716144000000,
  "principalUsers": "张三,李四",
  "planShareUrl": "/track/share-plan-report?shareId=abc123def456",
  "syncTime": 1716230400000,
  "traceId": "trace-20260424-7371-callback-001"
}'''
p = doc.add_paragraph()
run = p.add_run(sample2)
run.font.name = 'Courier New'
run.font.size = Pt(9)

# ========== 保存 ==========
output_path = '/Users/zhaozhiwei/IdeaProjects/metersphere/docs/全流程平台对接/全流程平台对接功能需求说明书.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
