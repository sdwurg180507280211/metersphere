#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成需求平台对接详细设计文档的Word版本
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def create_word_document():
    """创建Word文档"""
    doc = Document()

    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 添加标题
    title = doc.add_heading('需求平台对接详细设计文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加文档信息
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light Grid Accent 1'
    info_table.cell(0, 0).text = '文档版本'
    info_table.cell(0, 1).text = 'V1.0'
    info_table.cell(1, 0).text = '编制日期'
    info_table.cell(1, 1).text = '2026年3月12日'
    info_table.cell(2, 0).text = '项目名称'
    info_table.cell(2, 1).text = '需求平台与MeterSphere测试平台对接'

    doc.add_page_break()

    return doc

def add_section_1(doc):
    """添加第1章：引言"""
    doc.add_heading('1. 引言', 1)

    p = doc.add_paragraph()
    p.add_run('本文档旨在对需求平台与MeterSphere测试平台对接功能的详细设计进行全面而准确的描述，为开发人员在实现软件功能时提供指导和参考。详细的设计规范和流程将有助于保证软件的稳定性、可维护性和可扩展性。')

    p = doc.add_paragraph()
    p.add_run('本文档涵盖需求同步、需求池管理、测试计划创建、状态回传等核心功能的详细设计，包括数据库设计、接口设计、类设计等内容。')

def add_section_2(doc):
    """添加第2章：概述"""
    doc.add_heading('2. 概述', 1)

    doc.add_heading('2.1 项目背景', 2)
    p = doc.add_paragraph()
    p.add_run('全流程平台（需求平台）与MeterSphere测试平台需要建立数据对接通道，实现需求数据的自动同步和测试计划的关联管理。')

    doc.add_paragraph('核心诉求：', style='List Bullet')
    doc.add_paragraph('需求平台的需求数据需要自动同步到MeterSphere', style='List Bullet 2')
    doc.add_paragraph('测试人员在MeterSphere中可以查看需求信息并创建测试计划', style='List Bullet 2')
    doc.add_paragraph('测试计划的状态和报告需要回传给需求平台', style='List Bullet 2')

    doc.add_paragraph('业务价值：', style='List Bullet')
    doc.add_paragraph('减少手工录入，提高效率', style='List Bullet 2')
    doc.add_paragraph('打通需求与测试的数据链路', style='List Bullet 2')
    doc.add_paragraph('实现需求全生命周期的可追溯性', style='List Bullet 2')

    doc.add_heading('2.2 设计目标', 2)

    doc.add_paragraph('功能目标：', style='List Bullet')
    doc.add_paragraph('实现需求平台到MeterSphere的需求数据同步（通过RocketMQ）', style='List Number 2')
    doc.add_paragraph('实现MeterSphere需求池功能，支持需求展示、筛选、查询', style='List Number 2')
    doc.add_paragraph('实现从需求池创建测试计划的流程，一个需求对应一个测试计划', style='List Number 2')
    doc.add_paragraph('实现测试计划状态和报告链接回传到需求平台', style='List Number 2')
    doc.add_paragraph('实现幂等处理、异常处理、历史数据初始化', style='List Number 2')

    doc.add_paragraph('性能目标：', style='List Bullet')
    doc.add_paragraph('消息消费延迟 < 5秒', style='List Bullet 2')
    doc.add_paragraph('需求池列表查询响应时间 < 2秒', style='List Bullet 2')
    doc.add_paragraph('支持并发创建测试计划，保证数据一致性', style='List Bullet 2')

def add_section_3(doc):
    """添加第3章：架构设计"""
    doc.add_heading('3. 架构设计', 1)

    doc.add_heading('3.1 总体架构', 2)
    p = doc.add_paragraph()
    p.add_run('需求平台通过RocketMQ发送需求消息到MeterSphere，MeterSphere消费消息写入需求池。测试人员在需求池中创建测试计划，测试计划状态变化时通过RocketMQ回传给需求平台。')

    doc.add_heading('3.2 模块设计', 2)

    doc.add_paragraph('需求同步模块', style='Heading 3')
    doc.add_paragraph('职责：消费RocketMQ消息、解析需求数据、幂等处理、写入需求池')

    doc.add_paragraph('需求池模块', style='Heading 3')
    doc.add_paragraph('职责：需求池数据管理、需求列表查询筛选分页、需求详情展示、状态管理')

    doc.add_paragraph('测试计划创建模块', style='Heading 3')
    doc.add_paragraph('职责：从需求池创建测试计划、需求编号绑定、状态联动更新、并发控制')

    doc.add_paragraph('状态回传模块', style='Heading 3')
    doc.add_paragraph('职责：监听测试计划状态变化、生成回传消息、发送到RocketMQ、失败记录和重试')

    doc.add_heading('3.3 数据库设计', 2)

    # 需求池主表
    doc.add_paragraph('需求池主表（requirement_pool）', style='Heading 3')
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '字段名'
    hdr_cells[1].text = '类型'
    hdr_cells[2].text = '长度'
    hdr_cells[3].text = '必填'
    hdr_cells[4].text = '说明'

    fields = [
        ('id', 'VARCHAR', '50', '是', '主键'),
        ('dmp_num', 'VARCHAR', '100', '是', '需求编号（唯一索引）'),
        ('requirement_name', 'VARCHAR', '500', '是', '需求名称'),
        ('pool_status', 'VARCHAR', '20', '是', '状态：PENDING/CREATED/CANCELLED'),
        ('system_name', 'VARCHAR', '200', '否', '所属系统'),
        ('req_manager_name', 'VARCHAR', '100', '否', '需求负责人'),
        ('create_time', 'BIGINT', '-', '否', '需求提出时间'),
        ('linked_plan_id', 'VARCHAR', '50', '否', '关联的测试计划ID'),
        ('last_sync_time', 'BIGINT', '-', '是', '最后同步时间'),
        ('created_at', 'BIGINT', '-', '是', '创建时间'),
        ('updated_at', 'BIGINT', '-', '是', '更新时间'),
    ]

    for field in fields:
        row_cells = table.add_row().cells
        for i, value in enumerate(field):
            row_cells[i].text = value

    doc.add_paragraph()
    doc.add_paragraph('索引设计：')
    doc.add_paragraph('PRIMARY KEY: id', style='List Bullet 2')
    doc.add_paragraph('UNIQUE INDEX: uk_dmp_num ON dmp_num', style='List Bullet 2')
    doc.add_paragraph('INDEX: idx_pool_status ON pool_status', style='List Bullet 2')

    # 测试计划表扩展
    doc.add_paragraph('测试计划表扩展（test_plan）', style='Heading 3')
    doc.add_paragraph('新增字段：requirement_number (VARCHAR 100) - 需求编号（唯一索引）')
    doc.add_paragraph('新增索引：UNIQUE INDEX uk_requirement_number ON requirement_number')
    doc.add_paragraph('状态枚举扩展：新增 CANCELLED（已取消）')

def add_section_5(doc):
    """添加第5章：接口设计"""
    doc.add_heading('5. 接口设计', 1)

    doc.add_heading('5.1 需求池查询接口', 2)
    doc.add_paragraph('接口路径：GET /requirement-pool/list')
    doc.add_paragraph('功能：分页查询需求池列表，支持多条件筛选')

    doc.add_heading('5.2 需求详情接口', 2)
    doc.add_paragraph('接口路径：GET /requirement-pool/{dmpNum}')
    doc.add_paragraph('功能：根据需求编号查询需求详情')

    doc.add_heading('5.3 创建测试计划接口', 2)
    doc.add_paragraph('接口路径：POST /test-plan/create-from-requirement')
    doc.add_paragraph('功能：从需求池创建测试计划，自动绑定需求编号')
    doc.add_paragraph('异常情况：')
    doc.add_paragraph('需求不存在：返回404', style='List Bullet 2')
    doc.add_paragraph('需求已创建测试计划：返回400', style='List Bullet 2')
    doc.add_paragraph('需求已取消：返回400', style='List Bullet 2')

    doc.add_heading('5.4 RocketMQ消息格式', 2)
    doc.add_paragraph('需求同步消息（需求平台 → MeterSphere）', style='Heading 3')
    doc.add_paragraph('Topic: topic-requirement-to-metersphere')
    doc.add_paragraph('operationType枚举：CREATED（新建）、UPDATED（更新）、CANCELLED（取消）')

    doc.add_paragraph('状态回传消息（MeterSphere → 需求平台）', style='Heading 3')
    doc.add_paragraph('Topic: topic-metersphere-to-requirement')
    doc.add_paragraph('包含字段：dmpNum、planStatus、时间字段、负责人、报告链接等')

def add_section_7(doc):
    """添加第7章：安全设计"""
    doc.add_heading('7. 安全设计', 1)

    doc.add_heading('7.1 消息传输安全', 2)
    doc.add_paragraph('RocketMQ认证：配置AccessKey和SecretKey，启用ACL权限控制')
    doc.add_paragraph('消息加密：使用TLS/SSL加密通道')

    doc.add_heading('7.2 接口访问控制', 2)
    doc.add_paragraph('权限控制：需求池查询需要登录权限，创建测试计划需要相应权限')
    doc.add_paragraph('数据隔离：需求池全局可见，测试计划按项目隔离')

    doc.add_heading('7.3 数据审计', 2)
    doc.add_paragraph('操作日志：记录需求同步、测试计划创建、状态回传等关键操作')
    doc.add_paragraph('追踪机制：使用traceId贯穿整个流程，便于问题排查')

    doc.add_heading('7.4 异常处理', 2)
    doc.add_paragraph('幂等控制：基于dmpNum + eventTime进行幂等判断')
    doc.add_paragraph('并发控制：数据库唯一索引 + 乐观锁')
    doc.add_paragraph('失败重试：消息消费失败自动重试，回传失败记录日志支持手工重放')

if __name__ == '__main__':
    print("开始生成Word文档...")
    doc = create_word_document()

    print("添加第1章：引言")
    add_section_1(doc)

    print("添加第2章：概述")
    add_section_2(doc)

    print("添加第3章：架构设计")
    add_section_3(doc)

    print("添加第5章：接口设计")
    add_section_5(doc)

    print("添加第7章：安全设计")
    add_section_7(doc)

    # 保存文档
    output_path = '/Users/zhaozhiwei/IdeaProjects/metersphere/docs/全流程平台对接/需求平台对接详细设计文档.docx'
    doc.save(output_path)
    print(f"Word文档已生成：{output_path}")

