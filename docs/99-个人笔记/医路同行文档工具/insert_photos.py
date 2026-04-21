
import os
import random
import zipfile
import tempfile
import shutil
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from PIL import Image

# 配置
BASE_PATH = '/Users/edy/Desktop/百度网盘同步空间/其它/审核/20260417'
DOC_PATH = '/Users/edy/Desktop/百度网盘同步空间/其它/审核/20260417/医路同行-慢病规范化诊疗提升项目 总结报告.docx'
OUTPUT_PATH = '/Users/edy/Desktop/百度网盘同步空间/其它/审核/20260417/医路同行-慢病规范化诊疗提升项目 总结报告_含照片.docx'

FOLDERS = [
    '单人会议1-150',
    '单人会议301-450',
    '单人会议600-796',
    '多人会议 51-100场'
]

MEETINGS_PER_FOLDER = 5  # 每个文件夹取5个会议
MAX_WIDTH = Cm(16)    # 最大宽度16cm
MAX_HEIGHT = Cm(10)   # 最大高度10cm


def get_scaled_dimensions(img_path):
    """根据图片尺寸计算缩放后的宽度和高度"""
    try:
        with Image.open(img_path) as img:
            orig_width, orig_height = img.size

            # 转换为cm（假设96 DPI）
            # 1 inch = 2.54 cm, 96 DPI
            cm_per_pixel = 2.54 / 96
            width_cm = orig_width * cm_per_pixel
            height_cm = orig_height * cm_per_pixel

            # 计算缩放比例
            width_ratio = MAX_WIDTH.cm / width_cm
            height_ratio = MAX_HEIGHT.cm / height_cm
            ratio = min(width_ratio, height_ratio, 1.0)  # 不放大原图

            final_width = Cm(width_cm * ratio)
            final_height = Cm(height_cm * ratio)

            return final_width, final_height
    except Exception as e:
        print(f"  读取图片尺寸失败: {os.path.basename(img_path)}, 使用默认宽度")
        return MAX_WIDTH, None


def get_meeting_images_zip(zip_path):
    """从单个zip获取课前+课中图片，返回(会议名, [图片路径])"""
    temp_dir = tempfile.mkdtemp()
    meeting_name = os.path.basename(zip_path).replace('.zip', '')
    images = []

    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            keqian = None
            kezhong = None
            for f in zf.namelist():
                filename = os.path.basename(f)
                if filename in ['课前.jpg', '课前.png'] and keqian is None:
                    keqian = f
                elif filename in ['课中.jpg', '课中.png'] and kezhong is None:
                    kezhong = f
                if keqian and kezhong:
                    break

            if keqian:
                zf.extract(keqian, temp_dir)
                path = os.path.join(temp_dir, keqian)
                if os.path.exists(path) and os.path.getsize(path) > 0:
                    images.append(path)
            if kezhong:
                zf.extract(kezhong, temp_dir)
                path = os.path.join(temp_dir, kezhong)
                if os.path.exists(path) and os.path.getsize(path) > 0:
                    images.append(path)
    except Exception as e:
        print(f"  读取失败: {zip_path}")
        shutil.rmtree(temp_dir)
        return None, []

    if images:
        return meeting_name, images, temp_dir
    else:
        shutil.rmtree(temp_dir)
        return None, []


def get_meeting_images_folder(folder_path):
    """从单个子文件夹获取课前+课中图片"""
    meeting_name = os.path.basename(folder_path)
    images = []

    keqian = None
    kezhong = None
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file in ['课前.jpg', '课前.png'] and keqian is None:
                keqian = os.path.join(root, file)
            elif file in ['课中.jpg', '课中.png'] and kezhong is None:
                kezhong = os.path.join(root, file)
        if keqian and kezhong:
            break

    if keqian and os.path.getsize(keqian) > 0:
        images.append(keqian)
    if kezhong and os.path.getsize(kezhong) > 0:
        images.append(kezhong)

    if images:
        return meeting_name, images, None
    else:
        return None, []


def process_folder(folder_name):
    """处理一个文件夹，返回选中的会议列表"""
    folder_path = os.path.join(BASE_PATH, folder_name)
    if not os.path.exists(folder_path):
        return []

    zip_files = [f for f in os.listdir(folder_path) if f.endswith('.zip')]
    subfolders = [f for f in os.listdir(folder_path)
                  if os.path.isdir(os.path.join(folder_path, f)) and not f.startswith('.')]

    meetings = []

    if zip_files:
        # 处理ZIP方式
        random.shuffle(zip_files)
        for zip_file in zip_files:
            if len(meetings) >= MEETINGS_PER_FOLDER:
                break
            zip_path = os.path.join(folder_path, zip_file)
            name, images, temp_dir = get_meeting_images_zip(zip_path)
            if name and images:
                meetings.append((name, images, temp_dir))
                print(f"  [{folder_name}] {name}")
    else:
        # 处理子文件夹方式
        random.shuffle(subfolders)
        for subfolder in subfolders:
            if len(meetings) >= MEETINGS_PER_FOLDER:
                break
            subfolder_path = os.path.join(folder_path, subfolder)
            name, images, temp_dir = get_meeting_images_folder(subfolder_path)
            if name and images:
                meetings.append((name, images, temp_dir))
                print(f"  [{folder_name}] {name}")

    return meetings


def insert_images_to_doc():
    # 加载文档
    doc = Document(DOC_PATH)

    # 找到"会议照片"的位置
    photo_section_index = None
    for i, para in enumerate(doc.paragraphs):
        if '会议照片' in para.text:
            photo_section_index = i
            break

    if photo_section_index is None:
        print("未找到'会议照片'章节")
        return

    print(f"在段落 {photo_section_index} 找到'会议照片'")

    # 收集所有会议
    all_meetings = []
    temp_dirs = []

    print("\n=== 选择会议 ===")
    for folder in FOLDERS:
        meetings = process_folder(folder)
        all_meetings.extend(meetings)
        for m in meetings:
            if m[2]:  # temp_dir
                temp_dirs.append(m[2])

    print(f"\n总共选择了 {len(all_meetings)} 个会议")

    # 打乱会议顺序
    random.shuffle(all_meetings)

    # 收集所有图片（保持同一会议的课前+课中在一起）
    all_images = []
    print("\n=== 会议详情 ===")
    for i, (name, images, _) in enumerate(all_meetings, 1):
        print(f"  会议{i}: {name} -> {', '.join(os.path.basename(img) for img in images)}")
        all_images.extend(images)

    # 插入图片
    image_count = 0
    page_num = 0

    for i in range(0, len(all_images), 2):
        page_num += 1

        # 分页
        if i > 0:
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(WD_BREAK.PAGE)

        # 第一张 - 智能缩放
        if i < len(all_images):
            img_path = all_images[i]
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.space_after = Cm(1)
                run = p.add_run()

                final_width, final_height = get_scaled_dimensions(img_path)
                if final_height:
                    run.add_picture(img_path, width=final_width, height=final_height)
                else:
                    run.add_picture(img_path, width=final_width)

                image_count += 1
                print(f"  插入{image_count}: {os.path.basename(img_path)} "
                      f"({final_width.cm:.1f}cm x {final_height.cm if final_height else 'auto':.1f}cm)")
            except Exception as e:
                print(f"插入失败: {img_path}")
                import traceback
                traceback.print_exc()

        # 第二张 - 智能缩放
        if i + 1 < len(all_images):
            img_path = all_images[i + 1]
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()

                final_width, final_height = get_scaled_dimensions(img_path)
                if final_height:
                    run.add_picture(img_path, width=final_width, height=final_height)
                else:
                    run.add_picture(img_path, width=final_width)

                image_count += 1
                print(f"  插入{image_count}: {os.path.basename(img_path)} "
                      f"({final_width.cm:.1f}cm x {final_height.cm if final_height else 'auto':.1f}cm)")
            except Exception as e:
                print(f"插入失败: {img_path}")
                import traceback
                traceback.print_exc()

    # 保存
    doc.save(OUTPUT_PATH)
    print(f"\n✅ 完成!")
    print(f"文档: {OUTPUT_PATH}")
    print(f"共 {len(all_meetings)} 个会议, {image_count} 张图片, {page_num} 页")
    print(f"图片限制：宽度≤16cm, 高度≤10cm, 按比例缩放")

    # 清理
    for temp_dir in temp_dirs:
        try:
            shutil.rmtree(temp_dir)
        except:
            pass


if __name__ == '__main__':
    insert_images_to_doc()
