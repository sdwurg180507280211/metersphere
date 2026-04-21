
from docx import Document

doc_path = '/Users/edy/Desktop/百度网盘同步空间/其它/审核/20260417/医路同行-慢病规范化诊疗提升项目 总结报告.docx'

try:
    doc = Document(doc_path)
    print(f"文档段落数: {len(doc.paragraphs)}")
    print(f"文档节数: {len(doc.sections)}")
    print(f"文档表格数: {len(doc.tables)}")
    print("\n=== 文档段落内容 ===")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text[:100]}")
    print("\n=== 查找'项目照片'位置 ===")
    for i, para in enumerate(doc.paragraphs):
        if '项目照片' in para.text:
            print(f"在段落 {i} 找到: {para.text}")
except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()
